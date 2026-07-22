"""
简历导出为 .docx —— 干净的排版，便于直接下载/打印（意见 E-3）。

不依赖原始简历版式，而是用 python-docx 重建一份专业排版的中/英简历。
原始简历版式的像素级复刻留作后续优化，本期先做「能直接下载使用的 Word」。
"""

from pathlib import Path
from datetime import datetime


def _docx_available() -> bool:
    try:
        import docx  # noqa: F401
        return True
    except Exception:
        return False


def export_resume_docx(customized: dict, job: dict = None, lang: str = "zh",
                       output_dir: Path = None) -> Path:
    """
    把定制后的简历 dict 导出为 .docx，返回文件路径。
    customized: generator.generate_resume 返回的结构化简历（含 summary/experience/projects/skills/education）
    lang: 'zh' | 'en'，自动取对应字段（缺英文时回退中文）
    """
    if not _docx_available():
        raise RuntimeError(
            "未安装 python-docx，无法导出 Word。请运行 `pip install python-docx` 后重试。"
        )
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if output_dir is None:
        output_dir = Path(__file__).parent.parent / "output"
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    is_en = (lang == "en")

    def g(key, key_en=None):
        if is_en and key_en:
            v = customized.get(key_en, "")
            return v if v else customized.get(key, "")
        return customized.get(key, "")

    doc = Document()

    # 基础样式
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    name = g("name", "name_en") or (customized.get("experience", [{}]) or [{}])[0].get("company", "Resume")
    email = g("email", "email_en")
    phone = g("phone", "phone_en")
    location = g("location", "location_en")

    h = doc.add_heading(name, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    contact = doc.add_paragraph()
    contact_run = contact.add_run("  |  ".join([x for x in [email, phone, location] if x]))
    contact_run.font.size = Pt(10)
    contact_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # 摘要
    summary = g("summary", "summary_en")
    if summary:
        doc.add_heading("摘要 / Summary" if is_en else "个人摘要", level=1)
        doc.add_paragraph(summary)

    # 技能
    skills = g("skills", "skills_en")
    if skills:
        doc.add_heading("技能 / Skills" if is_en else "专业技能", level=1)
        if isinstance(skills, list):
            doc.add_paragraph("、 ".join(skills))
        else:
            doc.add_paragraph(str(skills))

    # 工作经历
    experience = customized.get("experience", []) or []
    if experience:
        doc.add_heading("工作经历 / Experience" if is_en else "工作经历", level=1)
        for exp in experience:
            title = exp.get("title_en" if is_en else "title", exp.get("title", ""))
            company = exp.get("company", "")
            duration = exp.get("duration", "")
            p = doc.add_paragraph()
            r = p.add_run(f"{title}  ·  {company}")
            r.bold = True
            if duration:
                d = doc.add_paragraph()
                dr = d.add_run(duration)
                dr.font.size = Pt(9)
                dr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            bullets = exp.get("bullets_en" if is_en else "bullets", exp.get("bullets", []))
            for b in bullets:
                doc.add_paragraph(str(b), style="List Bullet")

    # 项目
    projects = customized.get("projects", []) or []
    if projects:
        doc.add_heading("项目经历 / Projects" if is_en else "项目经历", level=1)
        for proj in projects:
            name_p = proj.get("name_en" if is_en else "name", proj.get("name", ""))
            p = doc.add_paragraph()
            r = p.add_run(name_p)
            r.bold = True
            desc = proj.get("description_en" if is_en else "description", proj.get("description", ""))
            if desc:
                doc.add_paragraph(desc)
            highlights = proj.get("highlights_en" if is_en else "highlights", proj.get("highlights", []))
            for hdesc in highlights:
                doc.add_paragraph(str(hdesc), style="List Bullet")

    # 教育
    education = customized.get("education", []) or []
    if education:
        doc.add_heading("教育背景 / Education" if is_en else "教育背景", level=1)
        for edu in education:
            school = edu.get("school", "")
            degree = edu.get("degree", "")
            major = edu.get("major", "")
            p = doc.add_paragraph()
            r = p.add_run(f"{school}  ·  {degree} in {major}")
            r.bold = True

    company = (job or {}).get("company", "company").replace(" ", "_")
    role = (job or {}).get("title", "role").replace(" ", "_")
    date_str = datetime.now().strftime("%Y%m%d")
    lang_suffix = "_en" if is_en else ""
    filename = f"resume_{company}_{role}{lang_suffix}_{date_str}.docx"
    filepath = output_dir / filename
    doc.save(str(filepath))
    return filepath


def export_cover_letter_docx(cover_letter: dict, job: dict = None,
                             output_dir: Path = None) -> Path:
    """把 Cover Letter dict 导出为 .docx（含中英双语正文）。"""
    if not _docx_available():
        raise RuntimeError(
            "未安装 python-docx，无法导出 Word。请运行 `pip install python-docx` 后重试。"
        )
    from docx import Document

    if output_dir is None:
        output_dir = Path(__file__).parent.parent / "output"
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    doc = Document()
    company = (job or {}).get("company", "company").replace(" ", "_")
    role = (job or {}).get("title", "role").replace(" ", "_")
    date_str = datetime.now().strftime("%Y%m%d")

    doc.add_heading("Cover Letter / 求职信", level=0)
    if cover_letter.get("subject"):
        doc.add_paragraph(cover_letter.get("subject", ""))
    if cover_letter.get("body"):
        doc.add_paragraph(cover_letter.get("body", ""))
    if cover_letter.get("body_en"):
        doc.add_paragraph("---")
        doc.add_paragraph(cover_letter.get("subject_en", "Cover Letter"))
        doc.add_paragraph(cover_letter.get("body_en", ""))

    filename = f"coverletter_{company}_{role}_{date_str}.docx"
    filepath = output_dir / filename
    doc.save(str(filepath))
    return filepath
